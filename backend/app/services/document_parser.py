"""标书文档解析器 —— 纯文本提取 + 视觉兜底。

参考 industrial-marketing-ai 项目的成熟方案：
- PDF: PyPDF2 按页提取文本（几秒跑完 148 页）
- DOCX: python-docx 直接读
- 智能判断：有文字的页直接用，没文字的扫描页转图片送 Vision API OCR
- 不依赖 RAG-Anything / mineru 做初始文本提取（保留给语义检索用）
"""

import logging
import os
import re
from typing import Any, Optional

from pypdf import PdfReader
from docx import Document as DocxDocument

from app.services.rag_service import get_vision_func

logger = logging.getLogger(__name__)

# 文本阈值：一页少于这个字符数就认为可能是扫描图片
TEXT_THRESHOLD = 80


async def extract_pdf_text(pdf_path: str) -> str:
    """从 PDF 提取全部文本（带页码标记）。

    Args:
        pdf_path: PDF 文件路径。

    Returns:
        带 [第N页] 标记的完整文本，可直接喂给 LLM。
    """
    reader = PdfReader(pdf_path)
    total = len(reader.pages)
    text_parts = []
    blank_count = 0

    for i, page in enumerate(reader.pages):
        page_text = (page.extract_text() or "").strip()
        page_text = _clean_text(page_text)
        if len(page_text) >= TEXT_THRESHOLD:
            text_parts.append(f"[第{i+1}页]\n{page_text}")
        else:
            blank_count += 1

    logger.info("PDF 文本提取: %d/%d 页有文字, %d 页空白", total - blank_count, total, blank_count)
    return "\n\n".join(text_parts)


async def extract_pdf_with_pages(pdf_path: str) -> list[dict[str, Any]]:
    """按页提取 PDF 内容，标记每页的文本量和解析模式。

    Args:
        pdf_path: PDF 文件路径。

    Returns:
        页面列表: [{page, content, char_count, parse_mode: "text"|"vision"}]
    """
    reader = PdfReader(pdf_path)
    total = len(reader.pages)
    pages = []

    for i, page in enumerate(reader.pages):
        page_text = (page.extract_text() or "").strip()
        page_text = _clean_text(page_text)
        char_count = len(page_text)

        if char_count >= TEXT_THRESHOLD:
            pages.append({
                "page": i + 1,
                "content": page_text,
                "char_count": char_count,
                "parse_mode": "text",
            })
        else:
            pages.append({
                "page": i + 1,
                "content": None,
                "char_count": char_count,
                "parse_mode": "vision",
            })

    text_pages = sum(1 for p in pages if p["parse_mode"] == "text")
    logger.info("PDF 页面分析: %d 页, %d 文本页, %d 需视觉 OCR", total, text_pages, total - text_pages)
    return pages


async def extract_docx_text(file_path: str) -> str:
    """从 DOCX 文件提取文本（含表格）。

    Args:
        file_path: DOCX 文件路径。

    Returns:
        完整文本。
    """
    doc = DocxDocument(file_path)
    text_parts = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            text_parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))

    return "\n\n".join(text_parts)


async def ocr_vision_pages(
    pdf_path: str, vision_pages: list[int], progress_cb=None
) -> dict[int, str]:
    """用 Vision API 对扫描页做 OCR。

    Args:
        pdf_path: PDF 文件路径。
        vision_pages: 需要 OCR 的页码列表（1-based）。
        progress_cb: 可选回调，每完成一页调用一次。

    Returns:
        {page_num: ocr_text} 字典。
    """
    if not vision_pages:
        return {}

    try:
        vision_func = await get_vision_func()
    except ValueError as e:
        logger.warning("Vision API 未配置，跳过扫描页 OCR: %s", e)
        return {}

    # 用 pdf2image 把指定页转成图片；Windows 上缺 Poppler 时降级到 pypdfium2
    convert_from_path = None
    try:
        from pdf2image import convert_from_path
    except ImportError:
        logger.warning("pdf2image 未安装，尝试使用 pypdfium2 转换扫描页")

    results: dict[int, str] = {}
    for page_num in vision_pages:
        try:
            import base64
            from io import BytesIO

            buf = BytesIO()
            if convert_from_path is not None:
                try:
                    images = convert_from_path(
                        pdf_path, dpi=200, first_page=page_num, last_page=page_num
                    )
                    if not images:
                        continue
                    images[0].save(buf, format="PNG")
                except Exception as e:
                    logger.warning("pdf2image 转换第 %d 页失败，尝试 pypdfium2: %s", page_num, str(e))
                    buf = BytesIO()
                    _render_pdf_page_with_pypdfium(pdf_path, page_num, buf)
            else:
                _render_pdf_page_with_pypdfium(pdf_path, page_num, buf)

            img_base64 = base64.b64encode(buf.getvalue()).decode()

            ocr_text = await vision_func(img_base64, "请提取图片中的所有文字内容，保持原有格式。")
            cleaned = _clean_text(ocr_text or "")
            results[page_num] = cleaned
            logger.info("Vision OCR: 第 %d 页, %d 字符", page_num, len(cleaned))

            if progress_cb:
                progress_cb(page_num)
        except Exception as e:
            logger.warning("Vision OCR 第 %d 页失败: %s", page_num, str(e))
            results[page_num] = ""

    return results


def _render_pdf_page_with_pypdfium(pdf_path: str, page_num: int, buf) -> None:
    """用 pypdfium2 渲染 PDF 指定页到 PNG buffer。"""
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(pdf_path)
    try:
        page = pdf[page_num - 1]
        bitmap = page.render(scale=2.5)
        image = bitmap.to_pil()
        image.save(buf, format="PNG")
    finally:
        pdf.close()


def _clean_text(text: str) -> str:
    """清洗文本：合并多余空白、去掉非打印字符。"""
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    # 去掉纯乱码行（非中英文数字标点的行）
    lines = []
    for line in text.split('\n'):
        stripped = line.strip()
        if not stripped:
            lines.append('')
        elif re.search(r'[一-鿿\w]', stripped):
            lines.append(stripped)
    return '\n'.join(lines)


def get_pdf_page_count(pdf_path: str) -> int:
    """快速获取 PDF 总页数。"""
    try:
        reader = PdfReader(pdf_path)
        return len(reader.pages)
    except Exception:
        return 0
