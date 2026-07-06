"""自动填写服务（P0-09）。

根据匹配结果，将符合条件的资质信息自动填入用户上传的投标文件模板。
支持 DOCX/PDF/XLSX 三种模板格式，输出 DOCX + PDF 双格式。
"""

import logging
import os
import re
import time
from datetime import datetime
from typing import Any, Optional

from app.config import OUTPUT_DIR, TEMPLATE_UPLOAD_DIR
from app.database import get_db
from app.models.template import FillTemplate
from app.services import file_convert
from app.utils.file_utils import detect_file_type, get_file_extension

logger = logging.getLogger(__name__)

# 占位符正则：{{字段名}}
_PLACEHOLDER_PATTERN = re.compile(r"\{\{([^}]+)\}\}")

# 字段映射：占位符名称 → 资质字段
_FIELD_MAP: dict[str, str] = {
    "证书名称": "name",
    "资质名称": "name",
    "证书编号": "number",
    "编号": "number",
    "发证日期": "issue_date",
    "有效期": "expiry_date",
    "有效期至": "expiry_date",
    "到期日期": "expiry_date",
    "发证机构": "issuing_authority",
    "认证范围": "scope",
    "范围": "scope",
    "等级": "level",
    "级别": "level",
    "持证主体": "holder",
    "企业名称": "holder",
    "持有人": "holder",
}


def _now_iso() -> str:
    """返回当前时间的 ISO 8601 字符串。"""
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


async def _get_matched_data(tender_id: int) -> dict[str, str]:
    """获取标书匹配到的资质数据，构建占位符映射。

    Args:
        tender_id: 标书 ID。

    Returns:
        占位符名称 → 值 的映射字典。
    """
    db = await get_db()
    try:
        # 获取所有 matched 状态的匹配结果
        cursor = await db.execute(
            "SELECT qualification_id FROM match_results WHERE tender_id = ? AND status = 'matched'",
            (tender_id,),
        )
        matched_rows = await cursor.fetchall()

        if not matched_rows:
            return {}

        # 获取匹配到的资质信息
        qual_ids = [r["qualification_id"] for r in matched_rows if r["qualification_id"]]
        if not qual_ids:
            return {}

        placeholders = ", ".join("?" * len(qual_ids))
        cursor = await db.execute(
            f"SELECT * FROM qualifications WHERE id IN ({placeholders})",
            qual_ids,
        )
        qual_rows = await cursor.fetchall()
    finally:
        await db.close()

    # 构建占位符映射
    field_values: dict[str, str] = {}
    for qual in qual_rows:
        qual_dict = dict(qual)
        for placeholder, field_name in _FIELD_MAP.items():
            value = qual_dict.get(field_name)
            if value and placeholder not in field_values:
                field_values[placeholder] = str(value)

    return field_values


def _fill_docx(template_path: str, output_path: str, field_values: dict[str, str]) -> tuple[int, int]:
    """填写 DOCX 模板中的占位符。

    使用 python-docx 遍历段落和表格，替换 {{字段名}} 占位符。

    Args:
        template_path: 模板文件路径。
        output_path: 输出文件路径。
        field_values: 占位符名称 → 值 的映射。

    Returns:
        (已填写字段数, 跳过字段数)。
    """
    from docx import Document

    doc = Document(template_path)
    filled_count = 0
    skipped_count = 0

    def replace_placeholders(text: str) -> tuple[str, int, int]:
        """替换文本中的占位符。

        Args:
            text: 原始文本。

        Returns:
            (替换后文本, 已填写数, 跳过数)。
        """
        filled = 0
        skipped = 0

        def replacer(match: re.Match) -> str:
            nonlocal filled, skipped
            placeholder_name = match.group(1).strip()
            if placeholder_name in field_values:
                filled += 1
                return field_values[placeholder_name]
            else:
                skipped += 1
                return match.group(0)  # 保留原占位符

        new_text = _PLACEHOLDER_PATTERN.sub(replacer, text)
        return new_text, filled, skipped

    # 遍历段落
    for paragraph in doc.paragraphs:
        new_text, f, s = replace_placeholders(paragraph.text)
        filled_count += f
        skipped_count += s
        if f > 0 or s > 0:
            # 需要更新段落文本（保留格式）
            for run in paragraph.runs:
                new_run_text, f2, s2 = replace_placeholders(run.text)
                run.text = new_run_text

    # 遍历表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    new_text, f, s = replace_placeholders(paragraph.text)
                    filled_count += f
                    skipped_count += s
                    if f > 0 or s > 0:
                        for run in paragraph.runs:
                            new_run_text, f2, s2 = replace_placeholders(run.text)
                            run.text = new_run_text

    doc.save(output_path)
    logger.info("DOCX 填写完成: filled=%d, skipped=%d, output=%s", filled_count, skipped_count, output_path)
    return filled_count, skipped_count


def _fill_xlsx(template_path: str, output_path: str, field_values: dict[str, str]) -> tuple[int, int]:
    """填写 XLSX 模板中的占位符。

    使用 openpyxl 遍历所有工作表的单元格，替换 {{字段名}} 占位符。

    Args:
        template_path: 模板文件路径。
        output_path: 输出文件路径。
        field_values: 占位符名称 → 值 的映射。

    Returns:
        (已填写字段数, 跳过字段数)。
    """
    from openpyxl import load_workbook

    wb = load_workbook(template_path)
    filled_count = 0
    skipped_count = 0

    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    def replacer(match: re.Match) -> str:
                        nonlocal filled_count, skipped_count
                        placeholder_name = match.group(1).strip()
                        if placeholder_name in field_values:
                            filled_count += 1
                            return field_values[placeholder_name]
                        else:
                            skipped_count += 1
                            return match.group(0)

                    new_value = _PLACEHOLDER_PATTERN.sub(replacer, cell.value)
                    cell.value = new_value

    wb.save(output_path)
    logger.info("XLSX 填写完成: filled=%d, skipped=%d, output=%s", filled_count, skipped_count, output_path)
    return filled_count, skipped_count


def _fill_pdf(template_path: str, output_path: str, field_values: dict[str, str]) -> tuple[int, int]:
    """填写 PDF 模板（转换为 DOCX 输出）。

    PDF 模板填写策略：用 pdfplumber 提取文本内容，识别占位符并替换，
    生成新的 DOCX 文档作为输出。

    Args:
        template_path: PDF 模板文件路径。
        output_path: 输出 DOCX 文件路径。
        field_values: 占位符名称 → 值 的映射。

    Returns:
        (已填写字段数, 跳过字段数)。
    """
    import pdfplumber
    from docx import Document

    filled_count = 0
    skipped_count = 0

    # 提取 PDF 文本
    extracted_lines: list[str] = []
    with pdfplumber.open(template_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                extracted_lines.extend(text.split("\n"))

    # 创建 DOCX 并填入替换后的文本
    doc = Document()

    for line in extracted_lines:
        def replacer(match: re.Match) -> str:
            nonlocal filled_count, skipped_count
            placeholder_name = match.group(1).strip()
            if placeholder_name in field_values:
                filled_count += 1
                return field_values[placeholder_name]
            else:
                skipped_count += 1
                return match.group(0)

        new_line = _PLACEHOLDER_PATTERN.sub(replacer, line)
        doc.add_paragraph(new_line)

    doc.save(output_path)
    logger.info("PDF→DOCX 填写完成: filled=%d, skipped=%d, output=%s", filled_count, skipped_count, output_path)
    return filled_count, skipped_count


async def upload_template(tender_id: int, file_name: str, file_content: bytes) -> FillTemplate:
    """保存上传的模板文件并创建 FillTemplate 记录。

    Args:
        tender_id: 标书 ID。
        file_name: 原始文件名。
        file_content: 文件二进制内容。

    Returns:
        创建的 FillTemplate 模型实例。

    Raises:
        ValueError: 文件格式不支持时抛出。
    """
    from app.config import SUPPORTED_TEMPLATE_FORMATS

    ext = get_file_extension(file_name)
    if ext not in SUPPORTED_TEMPLATE_FORMATS:
        raise ValueError(f"不支持的模板格式: {ext}，支持: {SUPPORTED_TEMPLATE_FORMATS}")

    # 保存文件
    timestamp = int(time.time())
    safe_name = os.path.basename(file_name).replace("..", "").replace("/", "").replace("\\", "")
    saved_name = f"{tender_id}_{timestamp}_{safe_name}"
    file_path = os.path.join(TEMPLATE_UPLOAD_DIR, saved_name)
    os.makedirs(TEMPLATE_UPLOAD_DIR, exist_ok=True)
    with open(file_path, "wb") as f:
        f.write(file_content)

    file_type = detect_file_type(file_path)
    now = _now_iso()

    db = await get_db()
    try:
        cursor = await db.execute(
            """INSERT INTO fill_templates
               (tender_id, filename, file_path, file_type, status, created_at)
               VALUES (?, ?, ?, ?, 'pending', ?)""",
            (tender_id, file_name, file_path, file_type, now),
        )
        await db.commit()
        template_id = cursor.lastrowid
    finally:
        await db.close()

    logger.info("模板上传成功: id=%s, tender_id=%s, filename=%s", template_id, tender_id, file_name)

    return FillTemplate(
        id=template_id,
        tender_id=tender_id,
        filename=file_name,
        file_path=file_path,
        file_type=file_type,
        status="pending",
        created_at=now,
    )


async def _update_template(template_id: int, **extra: Any) -> None:
    """更新模板记录。"""
    db = await get_db()
    try:
        sets = [f"{k} = ?" for k in extra]
        params = list(extra.values())
        params.append(template_id)
        await db.execute(
            f"UPDATE fill_templates SET {', '.join(sets)} WHERE id = ?",
            params,
        )
        await db.commit()
    finally:
        await db.close()


async def fill_template(tender_id: int) -> dict[str, Any]:
    """执行自动填写流程。

    流程：
    1. 获取标书的匹配结果（matched 的）
    2. 获取模板文件，识别格式
    3. 根据格式调用对应填写器
    4. 生成 DOCX 输出
    5. 用 file_convert 转 PDF（双输出）
    6. 返回填写结果

    Args:
        tender_id: 标书 ID。

    Returns:
        {"docx_path": ..., "pdf_path": ..., "filled_count": N, "skipped_count": M}

    Raises:
        ValueError: 模板不存在时抛出。
        RuntimeError: 填写失败时抛出。
    """
    logger.info("开始自动填写: tender_id=%s", tender_id)

    # 获取模板信息
    db = await get_db()
    try:
        cursor = await db.execute(
            "SELECT * FROM fill_templates WHERE tender_id = ? ORDER BY id DESC LIMIT 1",
            (tender_id,),
        )
        row = await cursor.fetchone()
    finally:
        await db.close()

    if row is None:
        raise ValueError(f"标书 {tender_id} 未上传模板，请先上传模板")

    template = FillTemplate(**dict(row))
    await _update_template(template.id or 0, status="filling")

    try:
        # 获取匹配的资质数据
        field_values = await _get_matched_data(tender_id)
        logger.info("获取到 %d 个占位符值: tender_id=%s", len(field_values), tender_id)

        # 确定输出路径
        timestamp = int(time.time())
        output_docx_path = os.path.join(OUTPUT_DIR, f"{tender_id}_filled_{timestamp}.docx")
        os.makedirs(OUTPUT_DIR, exist_ok=True)

        # 根据模板格式调用对应填写器
        file_type = template.file_type or detect_file_type(template.file_path or "")

        if file_type == "docx":
            filled_count, skipped_count = _fill_docx(
                template.file_path or "", output_docx_path, field_values
            )
        elif file_type == "xlsx":
            # XLSX 模板：先填写 XLSX，再转为 DOCX
            xlsx_output = os.path.join(OUTPUT_DIR, f"{tender_id}_filled_{timestamp}.xlsx")
            filled_count, skipped_count = _fill_xlsx(
                template.file_path or "", xlsx_output, field_values
            )
            # 将 XLSX 内容转为 DOCX
            from openpyxl import load_workbook
            from docx import Document
            doc = Document()
            wb = load_workbook(xlsx_output)
            for ws in wb.worksheets:
                doc.add_heading(f"工作表: {ws.title}", level=2)
                for row in ws.iter_rows(values_only=True):
                    row_text = " | ".join(str(c) if c is not None else "" for c in row)
                    doc.add_paragraph(row_text)
            doc.save(output_docx_path)
        elif file_type == "pdf":
            filled_count, skipped_count = _fill_pdf(
                template.file_path or "", output_docx_path, field_values
            )
        else:
            raise ValueError(f"不支持的模板格式: {file_type}")

        # 转换为 PDF（双输出）
        pdf_path: Optional[str] = None
        try:
            pdf_path = await file_convert.convert_to_pdf(output_docx_path, OUTPUT_DIR)
        except Exception as e:
            logger.warning("DOCX 转 PDF 失败（不影响 DOCX 输出）: %s", str(e))
            pdf_path = None

        # 更新模板记录
        await _update_template(
            template.id or 0,
            status="completed",
            filled_path=output_docx_path,
            output_pdf_path=pdf_path,
        )

        result = {
            "docx_path": output_docx_path,
            "pdf_path": pdf_path,
            "filled_count": filled_count,
            "skipped_count": skipped_count,
        }

        logger.info(
            "自动填写完成: tender_id=%s, filled=%d, skipped=%d",
            tender_id,
            filled_count,
            skipped_count,
        )
        return result

    except Exception as e:
        logger.error("自动填写失败: tender_id=%s, error=%s", tender_id, str(e), exc_info=True)
        await _update_template(template.id or 0, status="failed")
        raise RuntimeError(f"自动填写失败: {str(e)}") from e


async def download_filled(tender_id: int, format: str = "docx") -> Optional[tuple[str, str]]:
    """获取填写结果文件路径。

    Args:
        tender_id: 标书 ID。
        format: 下载格式（docx/pdf）。

    Returns:
        (文件路径, 文件名) 元组，不存在时返回 None。
    """
    db = await get_db()
    try:
        cursor = await db.execute(
            "SELECT * FROM fill_templates WHERE tender_id = ? AND status = 'completed' "
            "ORDER BY id DESC LIMIT 1",
            (tender_id,),
        )
        row = await cursor.fetchone()
    finally:
        await db.close()

    if row is None:
        return None

    template = FillTemplate(**dict(row))

    if format == "pdf":
        file_path = template.output_pdf_path
        if not file_path or not os.path.isfile(file_path):
            return None
        filename = f"{template.filename or 'filled'}.pdf"
    else:
        file_path = template.filled_path
        if not file_path or not os.path.isfile(file_path):
            return None
        filename = f"{template.filename or 'filled'}.docx"

    return file_path, filename
